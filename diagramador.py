"""
Diagramador de simulados (Word -> layout final)
================================================
Le um .docx "cru" (conteudo certo, sem diagramacao) e devolve um novo
documento construido a partir do TEMPLATE_BASE (assets/template_base.docx)
-- um arquivo que ja tem o cabecalho/rodape com a arte definitiva (caixa
flutuante com titulo/ementa + foto + icones), feito direto no Word.
Nunca recriamos cabecalho/rodape via codigo; so trocamos o texto do
titulo/ementa dentro da caixa flutuante do cabecalho e escrevemos o
corpo (blocos, questoes, alternativas, gabarito) do zero.

Padrao de estilo do corpo:
  - fonte Avenir Book, 12pt no texto corrido, 13.5pt cor #014391 nos titulos;
  - texto justificado; alternativas com o mesmo recuo do enunciado (sem recuo extra);
  - letras das alternativas sem negrito;
  - espacamento entre linhas multiplo de 1.16;
  - fundo azul claro SO na secao "Revisao de 24 horas".
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "Avenir Book"
FONT_SIZE = Pt(12)
TITLE_SIZE = Pt(13.5)
TITLE_COLOR = RGBColor(0x01, 0x43, 0x91)
LINE_SPACING = 1.16

AZUL_CAIXA = "D9EEFB"  # fundo azul claro da caixa de revisao (ajustavel)

ASSETS_DIR = os.path.join(os.path.dirname(__file__), "assets")
TEMPLATE_BASE = os.path.join(ASSETS_DIR, "template_base.docx")

QUESTAO_RE = re.compile(r"^Quest(?:ã|a)o\s+(\d+)\b\s*(.*)$", re.IGNORECASE)
ALT_RE = re.compile(r"^([A-E])\)\s*(.*)$")
GAB_COMENTADO_Q_RE = re.compile(
    r"Quest(?:ã|a)o\s+(\d+)\s*[—–\-]\s*Alternativa correta:\s*([A-E])", re.IGNORECASE
)
GAB_TITLE_RE = re.compile(
    r"^(Quest(?:ã|a)o\s+\d+\s*[—–\-]\s*Alternativa correta:\s*[A-E])\s*\*?\s*(.*)$",
    re.IGNORECASE,
)
BLOCO_RE = re.compile(r"^BLOCO\s+(\d+)", re.IGNORECASE)
GABARITO_HEAD_RE = re.compile(r"^GABARITO COMENTADO", re.IGNORECASE)
REVISAO_START_RE = re.compile(r"^Revis(?:ã|a)o de 24 horas$", re.IGNORECASE)
META_GENERICA_RE = re.compile(r"^Meta\s+\d+\s*[—–\-]", re.IGNORECASE)
META_REVISAO_RE = re.compile(r"^Meta de Revis(?:ã|a)o", re.IGNORECASE)
GABARITO_DIVIDER_RE = re.compile(r"^[━_\-\s]*GABARITO[━_\-\s]*$", re.IGNORECASE)
ALT_SPLIT_RE = re.compile(r"\s(?=[A-E]\)\s)")


WHITESPACE_RE = re.compile(r"[\t\n\r\v\f]+")
MULTISPACE_RE = re.compile(r" {2,}")


def normalize_text(text):
    """Remove tabs/quebras de linha embutidas no meio do texto (herdadas
    do arquivo cru) e colapsa espacos duplicados. Sem isso, um caractere
    de quebra "\\n" sobrevive dentro do run e o Word o renderiza como se
    fosse uma quebra manual (Shift+Enter), mesmo que no nosso codigo a
    gente so tenha criado paragrafos de verdade."""
    text = WHITESPACE_RE.sub(" ", text)
    text = MULTISPACE_RE.sub(" ", text)
    return text.strip()


def format_paragraph(p, justify=True, space_before=None, space_after=4):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def style_run(run, bold=False, title=False, color=None, size=None):
    run.font.name = FONT_NAME
    run.bold = bold
    if title:
        run.font.size = size or TITLE_SIZE
        run.font.color.rgb = color or TITLE_COLOR
    else:
        run.font.size = size or FONT_SIZE
        if color:
            run.font.color.rgb = color


def add_body_paragraph(container, text=None, justify=True, space_before=None, space_after=4):
    p = container.add_paragraph()
    format_paragraph(p, justify=justify, space_before=space_before, space_after=space_after)
    if text:
        r = p.add_run(text)
        style_run(r)
    return p


def add_title_paragraph(container, text, justify=False, align_center=False, space_before=None, space_after=6):
    p = container.add_paragraph()
    format_paragraph(p, justify=justify, space_before=space_before, space_after=space_after)
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    style_run(r, bold=True, title=True)
    return p


def set_textbox_paragraph_text(paragraph_el, new_text):
    """Troca o texto de um <w:p> mantendo a formatacao do 1o run
    (fonte, tamanho, cor, negrito) -- usado dentro das caixas flutuantes
    do cabecalho, que python-docx nao expoe via paragraphs normal."""
    runs = paragraph_el.findall(qn("w:r"))
    if not runs:
        return
    first_run = runs[0]
    t_elements = first_run.findall(qn("w:t"))
    if t_elements:
        t_elements[0].text = new_text
        for extra_t in t_elements[1:]:
            first_run.remove(extra_t)
    else:
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = new_text
        first_run.append(t)
    for r in runs[1:]:
        r.getparent().remove(r)


def update_header_title(doc, titulo, ementa):
    """Atualiza o titulo/ementa dentro da caixa flutuante azul do
    cabecalho do TEMPLATE_BASE. Existem 2 copias do mesmo texto no XML
    (uma para o Word moderno, outra de fallback para versoes antigas) --
    as duas precisam ser atualizadas para ficarem consistentes."""
    header_el = doc.sections[0].header._element
    boxes = header_el.findall(".//" + qn("w:txbxContent"))
    for box in boxes:
        paragraphs = box.findall(qn("w:p"))
        if len(paragraphs) >= 1:
            set_textbox_paragraph_text(paragraphs[0], titulo)
        if len(paragraphs) >= 2:
            set_textbox_paragraph_text(paragraphs[1], ementa)


def clear_body(doc):
    """Remove todos os paragrafos/tabelas do corpo, preservando section/header/footer."""
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def extract_gabarito(paragraphs):
    gabarito = {}
    for text in paragraphs:
        for m in GAB_COMENTADO_Q_RE.finditer(text):
            gabarito[int(m.group(1))] = m.group(2).upper()
    return gabarito


def set_table_no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def shade_cell(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=10, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.font.name = FONT_NAME
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color


def split_question(text):
    """Quando a questao e as alternativas A-E vem todas coladas num so
    paragrafo (sem quebra de linha), separa em (numero, enunciado, [(letra, texto), ...])."""
    m = QUESTAO_RE.match(text)
    if not m:
        return None
    numero, resto = m.groups()
    partes = ALT_SPLIT_RE.split(resto)
    enunciado = partes[0].strip()
    alternativas = []
    for chunk in partes[1:]:
        am = ALT_RE.match(chunk.strip())
        if am:
            alternativas.append(am.groups())
    return numero, enunciado, alternativas


def start_revisao_box(doc, titulo_box):
    """Cria a caixa de fundo azul claro da secao 'Revisao de 24 horas'
    -- a UNICA parte do documento com fundo azul -- e devolve a celula
    onde o conteudo deve ser escrito."""
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.width = Cm(17)
    shade_cell(cell, AZUL_CAIXA)

    p = cell.paragraphs[0]
    format_paragraph(p, justify=False, space_after=6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titulo_box)
    style_run(r, bold=True, title=True)
    return cell


def add_question_paragraph(container, numero, enunciado):
    titulo_p = add_body_paragraph(container, space_before=12, space_after=0, justify=False)
    r1 = titulo_p.add_run(f"Questão {numero}")
    style_run(r1, bold=True)

    enunciado_p = add_body_paragraph(container, space_before=None, space_after=4)
    r2 = enunciado_p.add_run(enunciado)
    style_run(r2)
    return enunciado_p


def add_alternativa_paragraph(container, letra, texto):
    ap = add_body_paragraph(container, space_after=2)
    r1 = ap.add_run(f"{letra}) ")
    style_run(r1, bold=False)
    r2 = ap.add_run(texto)
    style_run(r2, bold=False)
    return ap


def add_gabarito_comentado_line(container, text):
    """Trata uma linha dentro do Gabarito Comentado. O padrao no arquivo
    cru e: 'Questao N — Alternativa correta: X * A) <texto da alt A>'
    (titulo e a 1a alternativa colados no mesmo paragrafo); as demais
    alternativas (B-E) e a Fundamentacao Legal vem em paragrafos
    separados. Aqui sempre separamos titulo (negrito) de alternativa
    (linha propria)."""
    gt_m = GAB_TITLE_RE.match(text)
    if gt_m:
        titulo, resto = gt_m.groups()
        tp = add_body_paragraph(container, space_before=10, space_after=3)
        tr = tp.add_run(titulo)
        style_run(tr, bold=True)

        resto = resto.strip()
        if resto:
            am = ALT_RE.match(resto)
            if am:
                letra, alt_texto = am.groups()
                add_alternativa_paragraph(container, letra, alt_texto)
            else:
                add_body_paragraph(container, resto)
        return

    alt_m = ALT_RE.match(text)
    if alt_m:
        letra, alt_texto = alt_m.groups()
        add_alternativa_paragraph(container, letra, alt_texto)
        return

    add_body_paragraph(container, text)


def add_centered_title_line(container, text):
    add_title_paragraph(container, text, justify=False, align_center=True, space_after=6)


def add_question_block(container, numero, enunciado, alternativas):
    add_question_paragraph(container, numero, enunciado)
    for letra, texto in alternativas:
        add_alternativa_paragraph(container, letra, texto)


def add_gabarito_table(doc, gabarito, n_blocos=5):
    if not gabarito:
        return
    doc.add_page_break()
    h = add_title_paragraph(doc, "GABARITO SIMPLIFICADO", align_center=True, space_after=10)

    numeros = sorted(gabarito.keys())
    total = len(numeros)
    por_bloco = -(-total // n_blocos)
    chunks = [numeros[i * por_bloco:(i + 1) * por_bloco] for i in range(n_blocos)]

    outer = doc.add_table(rows=1, cols=n_blocos)
    outer.alignment = WD_TABLE_ALIGNMENT.CENTER
    outer.autofit = False
    set_table_no_borders(outer)

    for b in range(n_blocos):
        cell = outer.cell(0, b)
        cell.width = Cm(3.2)

        hp = cell.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(f"Bloco {b + 1}")
        style_run(hr, bold=True, title=True, size=Pt(11))

        chunk = chunks[b]
        if not chunk:
            continue
        nested = cell.add_table(rows=len(chunk), cols=2)
        nested.autofit = False
        set_table_no_borders(nested)
        for col in nested.columns:
            col.width = Cm(1.5)

        for i, numero in enumerate(chunk):
            c1, c2 = nested.cell(i, 0), nested.cell(i, 1)
            set_cell_text(c1, str(numero), bold=True, size=10)
            set_cell_text(c2, gabarito[numero], bold=True, size=10)
            if i % 2 == 1:
                shade_cell(c1, "EFEFEF")
                shade_cell(c2, "EFEFEF")

        # paragrafo apos a tabela aninhada (exigencia do formato de tabela)
        cell.add_paragraph()


def diagramar(input_stream, titulo=None, subtitulo=None):
    """Recebe um stream (BytesIO) do docx cru e devolve (Document, n_gabarito).

    O documento devolvido E O TEMPLATE_BASE (cabecalho/rodape com a arte
    definitiva, feita no Word) com o corpo preenchido a partir do
    conteudo do arquivo cru."""
    raw = Document(input_stream)

    all_texts = [normalize_text(p.text) for p in raw.paragraphs]
    non_empty = [t for t in all_texts if t]
    titulo = titulo or (non_empty[0] if len(non_empty) > 0 else "")
    subtitulo = subtitulo or (non_empty[1] if len(non_empty) > 1 else "")

    original_paragraph_texts = all_texts

    doc = Document(TEMPLATE_BASE)
    update_header_title(doc, titulo, subtitulo)

    gabarito = extract_gabarito(original_paragraph_texts)

    clear_body(doc)

    skip_first_lines = 2
    count_skipped = 0
    target = doc
    box_mode = None  # None | "revisao" | "gabarito"

    for text in original_paragraph_texts:
        text = text.strip()

        if count_skipped < skip_first_lines and text in (titulo, subtitulo):
            count_skipped += 1
            continue

        if not text:
            continue

        if REVISAO_START_RE.match(text):
            target = start_revisao_box(doc, text)
            box_mode = "revisao"
            continue

        if box_mode == "revisao" and GABARITO_DIVIDER_RE.match(text):
            target = doc
            box_mode = "gabarito"
            doc.add_page_break()
            add_title_paragraph(doc, "GABARITO COMENTADO")
            continue

        if box_mode == "revisao":
            if META_REVISAO_RE.match(text) or META_GENERICA_RE.match(text):
                add_centered_title_line(target, text)
                continue
            q = split_question(text)
            if q:
                add_question_block(target, *q)
                continue
            add_body_paragraph(target, text)
            continue

        bloco_m = BLOCO_RE.match(text)
        gab_head_m = GABARITO_HEAD_RE.match(text)

        if bloco_m:
            box_mode = None
            target = doc
            doc.add_page_break()
            add_title_paragraph(doc, text)
            continue

        if gab_head_m:
            box_mode = "gabarito"
            target = doc
            doc.add_page_break()
            add_title_paragraph(doc, text)
            continue

        if box_mode == "gabarito":
            add_gabarito_comentado_line(target, text)
            continue

        questao_m = QUESTAO_RE.match(text)
        alt_m = ALT_RE.match(text)

        if questao_m:
            numero, resto = questao_m.groups()
            add_question_paragraph(doc, numero, resto)
            continue

        if alt_m:
            letra, resto = alt_m.groups()
            add_alternativa_paragraph(doc, letra, resto)
            continue

        add_body_paragraph(doc, text)

    add_gabarito_table(doc, gabarito)

    return doc, len(gabarito)
